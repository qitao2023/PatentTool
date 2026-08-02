"""
审查意见通知书 DOCX 生成模块

把 Markdown 格式的通知书按 office_action 版式规范渲染为 DOCX（从零生成，不含模板文件）：
  - 页面 A4，页边距上下 2.54cm、左右 3.17cm；
  - 正文 宋体/方正书宋_GBK，12pt，两端对齐，首行缩进 2 字符，行距最小值 20 磅；
  - 标题段（#/##/### 或整段加粗）加粗；
  - | 表格 → Word 表格（Table Grid）；
  - > 引用 → 楷体或斜体段；
  - 化学式下标：将 _x_ / _{1-x} 等标记转为 run 下标（run.font.subscript）。
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


# 化学式下标标记：_x_ → 下标 x；_{1-x} → 下标 {1-x}；M_x Si_1-x O_2 兼容处理
_SUB_RE = re.compile(r"_\{([^}]+)\}|_([A-Za-z0-9]+)_")
_EAST_ASIA_FONT = "宋体"


def _set_east_asia_font(run, name: str = _EAST_ASIA_FONT):
    """同时设置 ascii 和 eastAsia 字体，保证中文字体生效。"""
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def _setup_page(doc: Document):
    """A4 页面，页边距上下 2.54cm、左右 3.17cm。"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def _style_body_paragraph(p):
    """应用 office_action 正文版式：12pt、两端对齐、首行缩进 2 字符、行距最小值 20 磅。"""
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    pf.line_spacing = Pt(20)
    # 首行缩进 2 字符（宋体 12pt 时约 24pt）
    pf.first_line_indent = Pt(24)


def _add_body_runs(paragraph, text: str, bold: bool = False):
    """把一段文本写入段落，处理化学式下标标记 _x_ 和加粗标记 **...**。"""
    text = text.strip()
    # 先去掉首尾 **（整段加粗由 bold 参数控制）
    if text.startswith("**") and text.endswith("**") and len(text) > 4:
        text = text[2:-2]
        bold = True

    # 用正则切分：普通文本 / 下标片段
    tokens = []
    last = 0
    for m in _SUB_RE.finditer(text):
        if m.start() > last:
            tokens.append(("text", text[last:m.start()]))
        sub = m.group(1) if m.group(1) is not None else m.group(2)
        tokens.append(("sub", sub))
        last = m.end()
    if last < len(text):
        tokens.append(("text", text[last:]))

    if not tokens:
        tokens = [("text", text)]

    for kind, value in tokens:
        run = paragraph.add_run(value)
        _set_east_asia_font(run)
        run.font.size = Pt(12)
        run.bold = bold
        if kind == "sub":
            run.font.subscript = True


def _flush_table(data, document):
    """渲染 markdown 表格。data: 逐行 cell 列表。"""
    if not data:
        return
    rows = len(data)
    cols = max(len(r) for r in data)

    # 优先用 'Table Grid' 样式；注意 add_table 在样式缺失时会先插入 tbl 再抛异常，
    # 因此必须先检测样式是否存在，避免产生空残表。
    try:
        style_names = {s.name for s in document.styles}
        has_grid = "Table Grid" in style_names
    except Exception:
        has_grid = False

    if has_grid:
        table = document.add_table(rows=rows, cols=cols, style="Table Grid")
    else:
        table = document.add_table(rows=rows, cols=cols)
        _add_table_borders(table)

    for i, row in enumerate(data):
        for j in range(cols):
            cell_text = row[j].strip() if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_body_runs(p, cell_text, bold=(i == 0))
    document.add_paragraph()


def _add_table_borders(table):
    """给表格添加全边框（Table Grid 样式缺失时使用）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblPr.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "auto",
        })
        borders.append(el)
    tblPr.append(borders)


def markdown_to_oa_docx(markdown: str, out_path: str | Path,
                        template_path: str | Path | None = None) -> Path:
    """把 Markdown 通知书渲染为 DOCX（从零生成，不依赖模板文件）。

    Args:
        markdown: 通知书 Markdown 文本
        out_path: 输出 docx 路径
        template_path: 已废弃，保留参数仅为兼容调用方；始终从零生成

    Returns:
        生成的 docx 路径
    """
    out_path = Path(out_path)

    doc = Document()
    _setup_page(doc)
    # Normal 样式默认字体
    try:
        st = doc.styles["Normal"]
        st.font.name = _EAST_ASIA_FONT
        st.font.size = Pt(12)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), _EAST_ASIA_FONT)
    except Exception:
        pass

    lines = markdown.split("\n")
    in_table = False
    table_data = []

    def flush():
        nonlocal table_data, in_table
        if table_data:
            _flush_table(table_data, doc)
            table_data = []
            in_table = False

    for line in lines:
        stripped = line.strip()

        # 表格
        if stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]:
            if "---" in stripped:
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_data.append(cells)
            in_table = True
            continue
        elif in_table:
            flush()

        # 空行
        if not stripped:
            flush()
            continue

        # 标题
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            _style_body_paragraph(p)
            _add_body_runs(p, stripped[4:], bold=True)
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            _style_body_paragraph(p)
            _add_body_runs(p, stripped[3:], bold=True)
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            _style_body_paragraph(p)
            _add_body_runs(p, stripped[2:], bold=True)
            continue

        # 引用
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            pf.line_spacing = Pt(20)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_body_runs(p, stripped[2:], bold=False)
            for run in p.runs:
                run.italic = True
            continue

        # 无序列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph()
            _style_body_paragraph(p)
            _add_body_runs(p, "· " + stripped[2:], bold=False)
            continue

        # 普通段落
        p = doc.add_paragraph()
        _style_body_paragraph(p)
        _add_body_runs(p, stripped, bold=False)

    flush()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
