"""
专利详情和对比分析报告面板
"""
from pathlib import Path

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTextBrowser, QPushButton, QLabel,
    QFileDialog, QMessageBox, QButtonGroup, QTabWidget, QStackedWidget,
)
from PySide6.QtCore import Slot
import markdown

from src.utils.patent_formatter import format_patent_text
from src.utils.patent_extract import extract_embodiments

_BROWSER_STYLE = """
    font-family: "Microsoft YaHei", "Segoe UI", sans-serif;
    font-size: 13px;
    background-color: #ffffff;
    border: 1px solid #d0d7de;
    border-radius: 4px;
    padding: 12px;
    line-height: 1.8;
"""


class ReportPanel(QWidget):
    """右下侧专利详情和对比分析报告"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._current_markdown: str = ""
        self._patent_data: dict | None = None
        self._ai_markdown: str = ""
        self._current_view: str = "detail"
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # ── 标题行 + 视图切换 ──────────────────────────────────────────
        title_row = QHBoxLayout()

        self.btn_detail = QPushButton("📋 专利详情")
        self.btn_detail.setCheckable(True)
        self.btn_detail.setChecked(True)
        self.btn_detail.setObjectName("viewToggleBtn")
        self.btn_detail.clicked.connect(lambda: self._switch_view("detail"))

        self.btn_ai = QPushButton("🤖 AI 分析")
        self.btn_ai.setCheckable(True)
        self.btn_ai.setObjectName("viewToggleBtn")
        self.btn_ai.clicked.connect(lambda: self._switch_view("ai"))

        title_row.addWidget(self.btn_detail)
        title_row.addWidget(self.btn_ai)
        title_row.addStretch(1)

        self.export_word_btn = QPushButton("📄 导出 Word")
        self.export_word_btn.setObjectName("exportBtn")
        self.export_word_btn.setEnabled(False)
        self.export_word_btn.clicked.connect(self._export_word)
        title_row.addWidget(self.export_word_btn)

        layout.addLayout(title_row)

        # ── AI 分析浏览器 ──────────────────────────────────────────────
        self.ai_browser = QTextBrowser()
        self.ai_browser.setOpenExternalLinks(True)
        self.ai_browser.setStyleSheet(f"QTextBrowser {{{_BROWSER_STYLE}}}")

        # ── 专利详情标签页 ──────────────────────────────────────────────
        self.detail_tabs = QTabWidget()
        self.detail_tabs.setStyleSheet("""
            QTabWidget::pane { border: 1px solid #d0d7de; }
            QTabBar::tab { padding: 6px 18px; font-size: 13px; }
        """)

        # Tab 1: 摘要
        self.tab_summary = QTextBrowser()
        self.tab_summary.setOpenExternalLinks(True)
        self.tab_summary.setStyleSheet(f"QTextBrowser {{{_BROWSER_STYLE}}}")
        self.detail_tabs.addTab(self.tab_summary, "📋 摘要")

        # Tab 2: 说明书
        self.tab_desc = QTextBrowser()
        self.tab_desc.setOpenExternalLinks(True)
        self.tab_desc.setStyleSheet(f"QTextBrowser {{{_BROWSER_STYLE}}}")
        self.detail_tabs.addTab(self.tab_desc, "📖 说明书")

        # Tab 3: 权利要求书
        self.tab_claims = QTextBrowser()
        self.tab_claims.setOpenExternalLinks(True)
        self.tab_claims.setStyleSheet(f"QTextBrowser {{{_BROWSER_STYLE}}}")
        self.detail_tabs.addTab(self.tab_claims, "📜 权利要求书")

        # Tab 4: 具体实施方式（下载时从说明书抽取的详述/实施例部分）
        self.tab_embodiments = QTextBrowser()
        self.tab_embodiments.setOpenExternalLinks(True)
        self.tab_embodiments.setStyleSheet(f"QTextBrowser {{{_BROWSER_STYLE}}}")
        self.detail_tabs.addTab(self.tab_embodiments, "🧪 具体实施方式")

        # ── 堆叠切换 ───────────────────────────────────────────────────
        self.stack = QStackedWidget()
        self.stack.addWidget(self.detail_tabs)   # index 0
        self.stack.addWidget(self.ai_browser)     # index 1
        layout.addWidget(self.stack, 1)

        # 按钮样式
        self.setStyleSheet("""
            QPushButton#viewToggleBtn {
                padding: 4px 14px; border: 1px solid #c0c0c0;
                border-radius: 3px; background: #f5f5f5;
            }
            QPushButton#viewToggleBtn:checked {
                background: #0078d4; color: white; border-color: #0078d4;
            }
        """)

        self._show_placeholder()

    # ── 视图切换 ────────────────────────────────────────────────────────

    def _switch_view(self, view: str):
        self._current_view = view
        self.btn_detail.setChecked(view == "detail")
        self.btn_ai.setChecked(view == "ai")

        if view == "detail":
            self.stack.setCurrentIndex(0)
            if self._patent_data:
                self._render_patent_detail()
        else:
            self.stack.setCurrentIndex(1)
            if self._ai_markdown:
                self._show_markdown(self._ai_markdown)
                self.export_word_btn.setEnabled(True)
            else:
                self.ai_browser.setHtml("""
                <div style="text-align:center;color:#8b949e;padding:60px 20px;">
                    <h2>等待 AI 分析...</h2>
                    <p>AI 分析结果尚未返回，请稍候或点击专利详情查看。</p>
                </div>
                """)

    # ── 专利详情 ────────────────────────────────────────────────────────

    def show_patent_detail(self, patent_data: dict):
        self._patent_data = patent_data
        if self._current_view == "detail":
            self._render_patent_detail()

    def _render_patent_detail(self):
        if not self._patent_data:
            return

        d = self._patent_data
        title = d.get("title", "")
        title = title.split(" - ", 1)[-1] if " - " in title else title

        # ── Tab 1: 摘要 + 基本信息 ────────────────────────────────────
        parts = [f'<h2>{self._esc(title)}</h2>']
        parts.append(self._build_meta_table(d))
        abstract = d.get("abstract") or d.get("abstract_snippet", "")
        if abstract:
            parts.append(f'<h3>摘要</h3><p>{self._esc(str(abstract))}</p>')
        else:
            parts.append('<p style="color:#888;">无摘要信息</p>')
        self.tab_summary.setHtml(self._wrap_html("".join(parts)))

        # ── Tab 1: 说明书 ────────────────────────────────────────────
        desc = d.get("description", "")
        if desc:
            html = format_patent_text(str(desc), "description")
            self.tab_desc.setHtml(html)
            self.detail_tabs.setTabVisible(1, True)
        else:
            self.tab_desc.setHtml(self._wrap_html(
                '<p style="color:#888;">无说明书信息</p>'))
            self.detail_tabs.setTabVisible(1, False)

        # ── Tab 2: 权利要求书 ────────────────────────────────────────
        claims = d.get("claims", "")
        if claims:
            html = format_patent_text(str(claims), "claims")
            self.tab_claims.setHtml(html)
            self.detail_tabs.setTabVisible(2, True)
        else:
            self.tab_claims.setHtml(self._wrap_html(
                '<p style="color:#888;">无权利要求信息</p>'))
            self.detail_tabs.setTabVisible(2, False)

        # ── Tab 3: 具体实施方式（优先下载存的字段，旧数据兜底现抽）──
        emb = str(d.get("embodiments") or "")
        if not emb and d.get("description"):
            emb = extract_embodiments(str(d.get("description", "")))
        if emb:
            self.tab_embodiments.setHtml(format_patent_text(emb, "description"))
            self.detail_tabs.setTabVisible(3, True)
        else:
            self.tab_embodiments.setHtml(self._wrap_html(
                '<p style="color:#888;">无具体实施方式信息</p>'))
            self.detail_tabs.setTabVisible(3, False)

        # 默认选中摘要
        self.detail_tabs.setCurrentIndex(0)
        self.btn_detail.setChecked(True)
        self.btn_ai.setChecked(False)
        self._current_view = "detail"
        self.stack.setCurrentIndex(0)

    def _build_meta_table(self, d: dict) -> str:
        rows = []
        for label, key in [
            ("公布号", "publication_number"),
            ("申请号", "application_number"),
            ("申请人", "applicant"),
            ("发明人", "inventor"),
            ("IPC", "ipc"),
            ("公开日", "publication_date"),
        ]:
            val = (d.get(key) or d.get("patent_number") if key == "publication_number" else "")
            if not val and key == "publication_number":
                val = d.get("patent_number", "")
            if val:
                rows.append(f'<tr><td><b>{label}</b></td><td>{self._esc(str(val))}</td></tr>')

        score = d.get("relevance_score") or d.get("fulltext_score", "")
        reason = d.get("relevance_reason") or d.get("fulltext_reason", "")
        if score:
            s = f'{score}'
            if reason:
                s += f' — {self._esc(str(reason))}'
            rows.append(f'<tr><td><b>相关度</b></td><td>{s}</td></tr>')

        if not rows:
            return ""
        return "<table>" + "".join(rows) + "</table>"

    # ── AI 分析 ──────────────────────────────────────────────────────────

    def _show_placeholder(self):
        self.stack.setCurrentIndex(1)
        self.ai_browser.setHtml("""
        <div style="text-align:center;color:#8b949e;padding:60px 20px;">
            <h2>等待分析...</h2>
            <p>完成检索和去重后，<b>点击左侧某篇专利</b>即可查看详情和对比分析。</p>
        </div>
        """)

    def show_loading(self, patent_number: str):
        import html as html_mod
        safe = html_mod.escape(str(patent_number))
        self.ai_browser.setHtml(f"""
        <div style="text-align:center;color:#8b949e;padding:60px 20px;">
            <h2>⏳ AI 正在分析...</h2>
            <p>正在对 <b>{safe}</b> 与本申请进行详细对比，请稍候...</p>
        </div>
        """)

    def show_single_comparison(self, patent_number: str, markdown_text: str):
        if not markdown_text or not markdown_text.strip():
            self.ai_browser.setHtml("<p>对比结果为空</p>")
            return
        self._current_markdown = markdown_text
        self._ai_markdown = markdown_text
        if self._current_view == "ai":
            self._show_markdown(markdown_text)
            self.export_word_btn.setEnabled(True)

    @Slot(object)
    def show_report(self, report):
        if hasattr(report, "html_content") and report.html_content:
            self.ai_browser.setHtml(report.html_content)
        elif hasattr(report, "markdown_content") and report.markdown_content:
            self._current_markdown = report.markdown_content
            self._ai_markdown = report.markdown_content
            self._show_markdown(report.markdown_content)
        else:
            self.ai_browser.setHtml("<p>报告内容为空</p>")
        self.export_word_btn.setEnabled(True)

    def _show_markdown(self, md: str):
        html = markdown.markdown(
            md, extensions=["tables", "fenced_code"],
            output_format="xhtml")
        styled = self._wrap_html(html)
        self.ai_browser.setHtml(styled)

    # ── 导出 ────────────────────────────────────────────────────────────

    def _export_word(self):
        if not self._current_markdown:
            QMessageBox.warning(self, "提示", "没有可导出的内容")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "导出 Word 报告", "对比分析报告.docx",
            "Word 文档 (*.docx)"
        )
        if not path:
            return

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            import re

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(11)

            lines = self._current_markdown.split('\n')
            in_table = False
            table_data = []

            def flush_table(data, document):
                if not data:
                    return
                rows = len(data)
                cols = max(len(r) for r in data)
                table = document.add_table(rows=rows, cols=cols, style='Table Grid')
                for i, row in enumerate(data):
                    for j, cell_text in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = cell_text.strip()
                        if i == 0:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
                document.add_paragraph()

            for line in lines:
                if line.startswith('|') and line.strip().endswith('|'):
                    if '---' in line:
                        continue
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    table_data.append(cells)
                    in_table = True
                    continue
                elif in_table:
                    flush_table(table_data, doc)
                    table_data = []
                    in_table = False

                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('**') and '**' in line[2:]:
                    p = doc.add_paragraph()
                    run = p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', line))
                    run.bold = True
                elif line.strip().startswith('- '):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                elif line.strip().startswith('> '):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip()[2:])
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()

            if table_data:
                flush_table(table_data, doc)

            doc.save(path)
            QMessageBox.information(self, "导出成功", f"报告已导出到:\n{path}")

        except Exception as e:
            QMessageBox.warning(self, "导出失败", f"无法生成 Word 文档:\n{e}")

    def reset(self):
        self._current_markdown = ""
        self._patent_data = None
        self._ai_markdown = ""
        self._current_view = "detail"
        self.btn_detail.setChecked(True)
        self.btn_ai.setChecked(False)
        self.export_word_btn.setEnabled(False)
        self._show_placeholder()

    # ── 工具 ────────────────────────────────────────────────────────────

    @staticmethod
    def _wrap_html(body: str) -> str:
        css = """
        body { font-family:"Microsoft YaHei",sans-serif; line-height:1.8;
               padding:10px; color:#222; }
        h1 { border-bottom:2px solid #333; padding-bottom:8px; }
        h2 { border-bottom:1px solid #999; padding-bottom:5px; margin-top:24px; }
        h3 { margin-top:18px; }
        table { border-collapse:collapse; width:100%; margin:8px 0 12px 0; }
        td { border:1px solid #ddd; padding:6px 10px; font-size:13px; }
        td:first-child { background:#f5f7fa; width:80px; white-space:nowrap; }
        th, td { border:1px solid #ccc; padding:8px; text-align:left; }
        th { background:#f0f0f0; }
        blockquote { border-left:3px solid #0078D4; padding-left:15px;
                     color:#555; background:#f5f9ff; }
        code { background:#f5f5f5; padding:1px 4px; border-radius:3px; }
        """
        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{css}</style></head>
<body>{body}</body></html>"""

    @staticmethod
    def _esc(text: str) -> str:
        import html as html_mod
        return html_mod.escape(str(text))
